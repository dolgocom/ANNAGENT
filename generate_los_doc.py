from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
    return row

def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

# =====================================================================
# COVER PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("LOS")
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = RGBColor(88, 56, 186)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Life Operating System")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(88, 56, 186)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("Персональная мультиагентная ИИ-система — цифровой штаб первого лица")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Версия: 1.0  |  Дата: {datetime.date.today().strftime('%d.%m.%Y')}  |  Конфиденциально")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# =====================================================================
# SECTION 1 — EXECUTIVE SUMMARY
# =====================================================================
add_heading("1. Executive Summary", 1, (88, 56, 186))

add_heading("1.1 Проблема", 2)
add_paragraph(
    "Владелец холдинга ежедневно принимает сотни решений, проводит десятки встреч и несёт персональную "
    "ответственность за критические процессы. В этом потоке отсутствует единый интеллектуальный центр, который:"
)
bullets = [
    "знает его физическое и ментальное состояние в режиме реального времени;",
    "приоритизирует события дня с учётом энергетики и здоровья;",
    "контролирует финансы, отношения, здоровье и коммуникации в одном месте;",
    "предупреждает о рисках до их наступления — а не фиксирует последствия."
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

add_paragraph(
    "\nРезультат — хроническая перегрузка, снижение качества стратегических решений, "
    "упущенные контакты и нарастающий дисбаланс между работой и жизнью."
)

add_heading("1.2 Решение — LOS (Life Operating System)", 2)
add_paragraph(
    "LOS — это персональная мультиагентная ИИ-система, работающая как цифровой штаб первого лица. "
    "Система состоит из 7 специализированных ИИ-агентов под управлением единого Master Orchestrator. "
    "Весь поток данных, решений и коммуникаций проходит через единый человеческий интерфейс — "
    "Chief of Staff (Аня) в Telegram. Система работает на русском языке."
)

add_heading("1.3 Ключевые принципы", 2)
principles = [
    ("Решение — всегда за человеком", "Система только рекомендует; любое действие — с явного одобрения CoS."),
    ("Проактивность", "Система предупреждает о риске до его наступления."),
    ("Контекст, а не данные", "Каждая рекомендация сопровождается объяснением причинно-следственной связи."),
    ("Конфиденциальность", "Чувствительные данные (аудио, биометрия) не передаются третьим сторонам."),
    ("Единый ритм дня", "Утренний брифинг 07:00 МСК + реактивный режим + вечерний дайджест 22:00 МСК."),
]
for title_p, desc in principles:
    p = doc.add_paragraph()
    run = p.add_run(f"• {title_p}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# =====================================================================
# SECTION 2 — ARCHITECTURE
# =====================================================================
add_heading("2. Общая архитектура системы", 1, (88, 56, 186))

add_heading("2.1 Концепция: 7 агентов + 1 Orchestrator", 2)
add_paragraph(
    "Система построена по принципу «звезда»: в центре — Master Orchestrator (центральный агент-координатор), "
    "вокруг него — 7 специализированных агентов, каждый из которых является экспертом своего домена."
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Компонент"
hdr[1].text = "Роль"
hdr[2].text = "Фаза"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, "5C38BA")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

agents_table = [
    ("Master Orchestrator", "Центральный координатор: маршрутизирует запросы, агрегирует ответы, управляет порядком вызовов и разрешает конфликты", "Core"),
    ("🔋 Neuro & Bio Agent", "Мониторинг физического и ментального ресурса; оптимизация расписания на основе Oura Ring", "✅ Фаза 1 MVP"),
    ("🎯 Decision Support Agent", "Агрегатор: мультилинзовый анализ решений; утренний (07:00) и вечерний (22:00) брифинги", "✅ Фаза 1 MVP"),
    ("🔮 Esoteric Intelligence Agent", "Астрологические, нумерологические и матрично-судьбовые прогнозы дня", "⏳ Фаза 2"),
    ("❤️ Health & Maintenance Agent", "Трекинг здоровья: анализы, препараты, врачебные визиты", "⏳ Фаза 2"),
    ("🤝 Network & Relationship Agent", "Управление деловыми и личными связями; коммуникационные ритмы", "⏳ Фаза 2"),
    ("💰 Treasury & Finance Agent", "Трекинг личных финансов через интеграцию с CoinKeeper", "⏳ Фаза 2"),
    ("🎙 Communication Intelligence Agent", "Психологический и юридический анализ переговоров (PLAUDE)", "⏳ Фаза 2"),
    ("👤 CoS — Аня", "Единственный человеческий интерфейс системы; Telegram-бот, русский язык", "—"),
]
for row_data in agents_table:
    add_table_row(table, row_data)

doc.add_paragraph()

add_heading("2.2 Потоки данных между агентами", 2)
flows = [
    "Oura Ring → Neuro & Bio → Decision Support",
    "CoinKeeper → Treasury → Decision Support",
    "PLAUDE (аудио + текст) → Communication Intel → Decision Support (при флаге high_risk)",
    "Network → Communication Intel (профили участников переговоров)",
    "Network → Decision Support (контекст по вовлечённым людям)",
    "Esoteric → Decision Support + Neuro & Bio (кросс-оценка качества дня)",
    "Health → Neuro & Bio (медицинский контекст) + Decision Support",
    "Neuro & Bio → Health (текущее состояние здоровья для контекста)",
    "Аня (CoS) → все агенты (ручные вводы данных)",
]
for f in flows:
    doc.add_paragraph(f, style='List Bullet')

add_heading("2.3 Правила приоритетов и конфликтов", 2)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = "Конфликт"
hdr2[1].text = "Правило"
for cell in hdr2:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

conflicts = [
    ("Neuro & Bio vs. Esoteric", "Физическое состояние имеет приоритет над эзотерической оценкой"),
    ("Network-напоминание + критическое истощение", "Оба сигнала передаются Ане одновременно; решение за ней"),
    ("Decision Support vs. Treasury (финансовый дедлайн)", "Флаг deadline_conflict с описанием рисков обоих сценариев"),
]
for row_data in conflicts:
    add_table_row(table2, row_data)

doc.add_page_break()

# =====================================================================
# SECTION 3 — AGENT DETAILS
# =====================================================================
add_heading("3. Детализация агентов", 1, (88, 56, 186))

# 3.1 Neuro & Bio
add_heading("3.1 Neuro & Bio Intelligence Agent  [✅ Фаза 1 MVP]", 2)
add_paragraph(
    "Основная цель: максимизировать продуктивность и work-life баланс через анализ физического, "
    "психологического и ментального состояния — с выдачей конкретных рекомендаций по расписанию и восстановлению."
)
add_paragraph("Источники данных:", bold=True)
sources_neuro = [
    "Oura Ring (авто): HRV, качество сна, ЧСС, Readiness Score",
    "Ручной ввод — Аня: Энергия, фокус, настроение (субъективная оценка)",
    "Ручной ввод — Аня: Тренировка (Да/Нет), Массаж (Да/Нет), Алкоголь (Да/Нет)",
    "Health & Maintenance Agent: Медицинский контекст",
    "Календарь (авто): Типы встреч, количество и временные слоты",
]
for s in sources_neuro:
    doc.add_paragraph(s, style='List Bullet')

add_paragraph("Типология встреч по нагрузке:", bold=True)
loads = [
    "🔴 Высокая: Стратегические, переговоры/сделки, публичные выступления, командировки",
    "🟡 Средняя: Операционные (команда), личные встречи 1-на-1",
    "🟢 Низкая: Административные, рутинные",
    "⚪ Защищённый ресурс: Восстановительные блоки (спорт, отдых)",
    "⚪ Неприкосновенно: Семья / личная жизнь — агент не трогает",
]
for l in loads:
    doc.add_paragraph(l, style='List Bullet')

add_paragraph("Пример вывода:", bold=True)
p = doc.add_paragraph()
p.add_run(
    "«Readiness Score сегодня 61 (4-й день снижения подряд). Стратегическая встреча с партнёрами в 14:00 несёт "
    "высокий когнитивный риск. Рекомендую перенести на среду 10:00–12:00. До встречи: 20-минутная прогулка "
    "без телефона. Семейный блок 18:00–20:00 защищён.»"
).italic = True

# 3.2 Decision Support
add_heading("3.2 Decision Support Agent  [✅ Фаза 1 MVP]", 2)
add_paragraph(
    "Основная цель: центральный мозг системы — агрегирует контекст всех 6 агентов, проводит "
    "мультилинзовый анализ решений, приоритизирует задачи дня, выступает арбитром при конфликтах агентов."
)
add_paragraph("5 линз анализа решений:", bold=True)
lenses = [
    ("Физическая", "Neuro & Bio: Readiness Score, паттерны усталости"),
    ("Финансовая", "Treasury: текущая финансовая картина"),
    ("Эзотерическая", "Esoteric: качество дня по трём системам"),
    ("Медицинская", "Health: актуальный медицинский контекст"),
    ("Репутационная", "Network: кто вовлечён, история отношений, риски"),
]
for lens, desc in lenses:
    p = doc.add_paragraph()
    p.add_run(f"• {lens}: ").bold = True
    p.add_run(desc)

add_paragraph("Триггеры:", bold=True)
triggers = [
    "07:00 МСК (ежедневно): автоматический утренний брифинг",
    "22:00 МСК (ежедневно): вечерний дайджест",
    "Личный запрос Ани: описывает ситуацию — агент собирает контекст и выдаёт анализ",
    "Флаг Treasury: обнаружено крупное финансовое движение",
    "Конфликт агентов: противоречивые рекомендации — агент выступает арбитром",
]
for t in triggers:
    doc.add_paragraph(t, style='List Bullet')

add_paragraph("Пример утреннего брифинга:", bold=True)
p = doc.add_paragraph()
p.add_run(
    "«07:00 МСК. Состояние: Readiness 68, 3-й день снижения. Качество дня (Esoteric): высокое — рекомендован "
    "для стратегических решений. Финансы: норма. Приоритеты: 1) Решение по сделке с Петровым (дедлайн 2 дня); "
    "2) Уточнить юридические гарантии до подписания. Рекомендация: перенести операционные встречи на послеобед; "
    "пик когнитивной работы — до 13:00. Именинники сегодня: нет.»"
).italic = True

# 3.3 Esoteric
add_heading("3.3 Esoteric Intelligence Agent  [⏳ Фаза 2]", 2)
add_paragraph(
    "Основная цель: комплексный эзотерический советник, ежедневно оценивающий качество дня по трём системам "
    "и предоставляющий рекомендации по срокам и характеру ключевых действий."
)
systems = [
    ("Астрология (транзиты)", "Благоприятные / сложные периоды для сделок, переговоров, запусков. По натальной карте."),
    ("Матрица судьбы", "Личные энергетические циклы, периоды силы и восстановления."),
    ("Нумерология", "Личный год, личный месяц, личный день — цифровые циклы влияния."),
]
for sys, desc in systems:
    p = doc.add_paragraph()
    p.add_run(f"• {sys}: ").bold = True
    p.add_run(desc)
add_paragraph("Входные данные: дата, время и место рождения (астрология, Матрица судьбы); имя (нумерология).")

# 3.4 Health
add_heading("3.4 Health & Maintenance Agent  [⏳ Фаза 2]", 2)
add_paragraph(
    "Основная цель: трекинг физического здоровья — динамика анализов, врачебные визиты, "
    "медикаменты и ресёрч назначенных препаратов."
)
funcs_health = [
    "Динамика показателей: выявляет устойчивые отклонения в анализах",
    "Ресёрч препаратов: при новом назначении — состав, побочные эффекты, взаимодействия",
    "Напоминания о визитах: плановые обследования и интервалы",
    "Мониторинг приёма: фиксирует факт приёма медикаментов / добавок",
]
for f in funcs_health:
    doc.add_paragraph(f, style='List Bullet')

# 3.5 Network
add_heading("3.5 Network & Relationship Agent  [⏳ Фаза 2]", 2)
add_paragraph(
    "Основная цель: управление деловыми и личными связями — единая карта отношений с приоритетами, "
    "ритмами касаний, напоминаниями о датах и персонализированной коммуникационной поддержкой."
)
add_paragraph("Круги отношений:", bold=True)
circles = [
    ("🔴 Ядро", "Семья, ближайшие партнёры", "Постоянно"),
    ("🟡 Близкий", "Ключевые деловые партнёры, близкие друзья", "2–4 раза в месяц"),
    ("🟢 Рабочий", "Команда, подрядчики, клиенты", "По необходимости"),
    ("⚪ Расширенный", "Знакомые, нетворк", "Раз в квартал"),
]
table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = "Круг"
hdr3[1].text = "Состав"
hdr3[2].text = "Ритм касания"
for row_data in circles:
    add_table_row(table3, row_data)
doc.add_paragraph()
add_paragraph("Ключевые функции:", bold=True)
funcs_net = [
    "Подбор подарков на основе интересов; хранит историю — исключает повторы",
    "Персональные поздравления: под каждого человека, не повторяет тон и образы. Только русский язык.",
    "Динамика отношений: фиксирует статус; флаг при долгом отсутствии касания",
    "Система напоминаний: дни рождения (за 7, 2 дня и в день), праздники, годовщины",
]
for f in funcs_net:
    doc.add_paragraph(f, style='List Bullet')

# 3.6 Treasury
add_heading("3.6 Treasury & Personal Finance Agent  [⏳ Фаза 2]", 2)
add_paragraph(
    "Основная цель: трекинг личных финансов — доходы, расходы, бюджет, динамика по месяцам, "
    "контроль регулярных платежей. Источник данных: CoinKeeper API."
)
funcs_treasury = [
    "Трекинг по категориям: расходы/доходы с разбивкой и контролем лимитов",
    "Алерты превышения бюджета: мгновенный флаг при выходе за лимит категории",
    "Аномалии: выявляет нестандартные расходы; запрашивает подтверждение у Ани",
    "Регулярные платежи: напоминания до дедлайна; флаг payment_missed при просрочке",
    "Критический баланс: алерт critical_balance при уходе баланса в минус",
]
for f in funcs_treasury:
    doc.add_paragraph(f, style='List Bullet')

# 3.7 Communication Intel
add_heading("3.7 Communication Intelligence Agent  [⏳ Фаза 2]", 2)
add_paragraph(
    "Основная цель: психологический и юридический анализ переговоров на основе аудиозаписей "
    "и транскрипций из PLAUDE — выявление манипуляций, скрытых конфликтов, профайлинг участников."
)
add_paragraph("Аналитические модули:", bold=True)
modules = [
    "Манипуляции: газлайтинг, давление, лесть, уклонение, искусственный дефицит времени",
    "Скрытые конфликты: противоречия между словами, избегаемые темы",
    "Верификация искренности: несоответствия в логике аргументов (признаки, не доказательства)",
    "Профайлинг личности: тип личности, мотивации, вероятные слабые точки",
    "Юридический анализ: риски позиции, уязвимости, рекомендации по следующим шагам",
]
for m in modules:
    doc.add_paragraph(m, style='List Bullet')
add_paragraph("Политика хранения: хранятся результаты анализа и краткие сведения об участниках. "
              "НЕ хранятся: аудиозаписи, полные транскрипции, чувствительные детали переговоров.")

doc.add_page_break()

# =====================================================================
# SECTION 4 — MEMORY ARCHITECTURE
# =====================================================================
add_heading("4. Архитектура памяти агента", 1, (88, 56, 186))

add_paragraph(
    "На основе анализа лучших практик 2025–2026 (CoALA Framework, Princeton/CMU) выбрана "
    "четырёхслойная архитектура памяти в сочетании с фреймворком Mem0 (open-source, Apache-2.0)."
)

add_heading("4.1 Четыре слоя памяти", 2)
table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
headers4 = ["Слой", "Что хранит", "Где хранится", "Время жизни"]
for i, h in enumerate(headers4):
    hdr4[i].text = h
    for para in hdr4[i].paragraphs:
        for run in para.runs:
            run.bold = True

layers = [
    ("⚡ Рабочая (Working)", "Текущий разговор, активный контекст", "LLM context window / память в RAM", "Пока идёт разговор"),
    ("📖 Эпизодическая (Episodic)", "Что происходило: «Аня в июле летела в Бали»", "PostgreSQL + векторные эмбеддинги", "Навсегда"),
    ("💡 Семантическая (Semantic)", "Факты о боссе, людях, предпочтениях", "Mem0 (pgvector + граф)", "Навсегда"),
    ("⚙️ Процедурная (Procedural)", "Как делать: промты, навыки агентов", "Код / System prompt", "Постоянно (в коде)"),
]
for row in layers:
    add_table_row(table4, row)
doc.add_paragraph()

add_heading("4.2 Mem0 — выбранный фреймворк", 2)
add_paragraph(
    "Mem0 — open-source REST API, который автоматически извлекает факты, предпочтения и отношения из разговоров "
    "и сохраняет их с возможностью векторного поиска. Лицензия: Apache 2.0. "
    "Работает поверх PostgreSQL с расширением pgvector."
)
add_paragraph("Три внутренних хранилища Mem0:", bold=True)
stores = [
    ("Vector Store (pgvector)", "Семантический поиск по смыслу — находит похожий контекст"),
    ("Graph Store", "Хранит связи: «Пётр — деловой партнёр», «Иван — муж сестры»"),
    ("KV Store", "Быстрые факты: имя, город, язык, предпочтения"),
]
for store, desc in stores:
    p = doc.add_paragraph()
    p.add_run(f"• {store}: ").bold = True
    p.add_run(desc)

add_paragraph("\nКак это работает:", bold=True)
add_paragraph(
    "Когда Аня говорит «Он предпочитает встречи до 12:00», Mem0 автоматически: "
    "1) извлекает факт; 2) сохраняет в Vector Store + KV Store; "
    "3) при следующем запросе возвращает контекст оркестратору → персонализированный ответ."
)

add_paragraph("Сравнение с альтернативами:", bold=True)
table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
hdr5[0].text = "Фреймворк"
hdr5[1].text = "Преимущества"
hdr5[2].text = "Недостатки"
for cell in hdr5:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

comparisons = [
    ("✅ Mem0 (выбран)", "Open-source, pgvector, автоизвлечение фактов, Python-библиотека, идеален для личных ассистентов", "Менее мощный граф чем Zep"),
    ("⚠️ Zep / Graphiti", "Мощный граф связей, темпоральность", "Требует Neo4j, сложнее в настройке, корпоративное решение"),
    ("❌ Letta / MemGPT", "OS-подобная иерархия памяти, исследовательски интересен", "Требует отдельный сервер, медленнее в продакшне"),
]
for row in comparisons:
    add_table_row(table5, row)

doc.add_page_break()

# =====================================================================
# SECTION 5 — USER FLOW / DAILY RHYTHM
# =====================================================================
add_heading("5. Ритм работы с системой", 1, (88, 56, 186))

add_heading("5.1 07:00 МСК — Утренний брифинг (автоматически)", 2)
add_paragraph("Orchestrator автоматически запускает цепочку агентов и формирует единое сообщение для Ани:")
briefing_blocks = [
    ("Состояние", "Readiness Score, субъективные ощущения (если введены), тренды за 7 дней"),
    ("Качество дня", "Эзотерическая оценка: общая, по каждой системе, с объяснением"),
    ("Расписание", "Тип и нагрузка встреч; флаги перегруза; предложения по перестановкам"),
    ("Важные даты", "Именинники, религиозные и государственные праздники контактов"),
    ("Финансы", "Флаги: превышения бюджета, ближайшие платежи"),
    ("Приоритеты дня", "Топ-задачи от Decision Support с контекстом и рекомендуемым тайм-слотом"),
]
table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
hdr6[0].text = "Блок"
hdr6[1].text = "Содержание"
for cell in hdr6:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
for row in briefing_blocks:
    add_table_row(table6, row)
doc.add_paragraph()

add_paragraph(
    "После брифинга Аня вводит субъективный статус: энергия / фокус / настроение / тренировка / массаж / алкоголь. "
    "Логика напоминаний: первый пинг +5 мин → если нет ответа, второй пинг +5 мин → "
    "после этого система работает на данных Oura без субъективного ввода."
)

add_heading("5.2 День — Реактивный режим", 2)
reactive_events = [
    "Новая встреча в календаре → Neuro & Bio оценивает нагрузку, при необходимости — рекомендация",
    "Новые данные из CoinKeeper → Treasury обновляет картину; мгновенный алерт при превышении бюджета",
    "Аня загружает запись из PLAUDE → Communication Intelligence проводит анализ переговоров",
    "Аня описывает дилемму → Decision Support собирает контекст и выдаёт мультилинзовый анализ",
    "Тревожный тренд в Oura → Health Agent проверяет медицинский контекст и флагирует",
]
for e in reactive_events:
    doc.add_paragraph(e, style='List Bullet')

add_heading("5.3 22:00 МСК — Вечерний дайджест (автоматически)", 2)
digest_blocks = [
    ("Итог дня", "Динамика Readiness / состояния: как прошёл день"),
    ("Финансы", "Что потрачено сегодня; отклонения от нормы"),
    ("Контакты", "Кто выпал из касания; пропущенные важные события"),
    ("Открытые задачи", "Что не закрыто из приоритетов дня"),
    ("Завтра", "Предварительная оценка следующего дня (состояние + Esoteric)"),
]
table7 = doc.add_table(rows=1, cols=2)
table7.style = 'Table Grid'
hdr7 = table7.rows[0].cells
hdr7[0].text = "Блок"
hdr7[1].text = "Содержание"
for cell in hdr7:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
for row in digest_blocks:
    add_table_row(table7, row)

doc.add_page_break()

# =====================================================================
# SECTION 6 — TECH STACK
# =====================================================================
add_heading("6. Технологический стек", 1, (88, 56, 186))

table8 = doc.add_table(rows=1, cols=3)
table8.style = 'Table Grid'
hdr8 = table8.rows[0].cells
hdr8[0].text = "Компонент"
hdr8[1].text = "Технология"
hdr8[2].text = "Назначение"
for cell in hdr8:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

tech_stack = [
    ("Telegram Bot", "Python + aiogram 3.x", "Единственный интерфейс пользователя"),
    ("LLM / Оркестратор", "OpenAI GPT-4o", "Принятие решений, генерация ответов, ReAct-паттерн"),
    ("База данных", "PostgreSQL + pgvector", "Долгосрочное хранение всех данных"),
    ("Память агента", "Mem0 (open-source)", "Автоизвлечение фактов, семантический поиск"),
    ("Веб-поиск", "Tavily Search API", "Агент поиска по интернету"),
    ("Биометрия", "Oura Ring API", "HRV, Readiness Score, сон, ЧСС"),
    ("Финансы", "CoinKeeper API / парсинг", "Доходы, расходы, бюджет"),
    ("Переговоры", "PLAUDE (аудио + транскрипция)", "Анализ встреч и переговоров"),
    ("Инфраструктура", "Replit (хостинг)", "Деплой и запуск системы"),
    ("Паттерн агентов", "ReAct (Reasoning + Acting)", "Думает → Вызывает инструмент → Анализирует результат"),
]
for row in tech_stack:
    add_table_row(table8, row)

doc.add_paragraph()

add_heading("6.1 Внешние API и ключи доступа (необходимы для запуска)", 2)
apis = [
    ("Telegram Bot Token", "Через @BotFather в Telegram (/newbot)"),
    ("OpenAI API Key", "platform.openai.com → API Keys"),
    ("Oura Ring Personal Access Token", "cloud.ouraring.com/personal-access-tokens"),
    ("Tavily API Key (поиск)", "tavily.com → бесплатный тариф доступен"),
    ("CoinKeeper API", "Уточняется: API или парсинг — на этапе разработки Фазы 2"),
]
table9 = doc.add_table(rows=1, cols=2)
table9.style = 'Table Grid'
hdr9 = table9.rows[0].cells
hdr9[0].text = "Сервис"
hdr9[1].text = "Как получить"
for cell in hdr9:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
for row in apis:
    add_table_row(table9, row)

doc.add_page_break()

# =====================================================================
# SECTION 7 — ROADMAP
# =====================================================================
add_heading("7. Дорожная карта реализации", 1, (88, 56, 186))

add_heading("Фаза 1 — MVP  ✅", 2)
add_paragraph("Состав: Neuro & Bio Intelligence Agent + Decision Support Agent + Telegram-бот.")
phase1 = [
    "Telegram-бот на русском языке — единственный интерфейс",
    "Интеграция с Oura Ring API (HRV, Readiness Score, сон, ЧСС)",
    "Ручной ввод субъективного состояния через бот (энергия, фокус, настроение, тренировка, массаж, алкоголь)",
    "Анализ календаря: типология встреч, когнитивная нагрузка дня",
    "Алерты при низком Readiness Score + рекомендации по перестановке в расписании",
    "Утренний брифинг в 07:00 МСК (автоматически)",
    "Вечерний дайджест в 22:00 МСК (автоматически)",
    "Базовая логика Decision Support: приоритизация задач дня",
    "Архитектура памяти: рабочая + эпизодическая (Mem0 + PostgreSQL)",
    "Двухступенчатая логика напоминаний при отсутствии ввода (+5 мин → +5 мин)",
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("Результат Фазы 1: ")
run.bold = True
p.add_run("Работающий MVP — владелец получает персонализированный утренний брифинг и рекомендации "
          "по расписанию на основе физического состояния.")

doc.add_paragraph()
add_heading("Фаза 2 — Полная интеграция  ⏳", 2)
add_paragraph("Состав: все 7 агентов + полная интеграция источников данных.")
phase2 = [
    "Esoteric Intelligence Agent (требует ввода натальных данных)",
    "Health & Maintenance Agent: трекинг анализов, визитов, препаратов; ресёрч назначений",
    "Network & Relationship Agent: карта контактов, напоминания, подарки, поздравления",
    "Treasury & Finance Agent: интеграция с CoinKeeper",
    "Communication Intelligence Agent: анализ PLAUDE (аудио + транскрипция)",
    "Полная кросс-агентная интеграция: потоки данных между всеми агентами",
    "Механизм арбитража конфликтов через Decision Support",
    "Адаптация брифингов к часовому поясу при командировках",
    "Полный четырёхслойный Mem0: Vector + Graph + KV + Episodic store",
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# =====================================================================
# SECTION 8 — SECURITY & RULES
# =====================================================================
add_heading("8. Принципы безопасности и ограничения", 1, (88, 56, 186))

security_rules = [
    "Биометрические данные (Oura) не передаются агентам без явной необходимости",
    "Аудиозаписи переговоров не хранятся после завершения анализа",
    "Данные банковских карт и реквизиты счётов не хранятся — только агрегированные данные из CoinKeeper",
    "Вывод рекомендаций идёт исключительно через Аню; система не контактирует с третьими лицами напрямую",
    "Система не принимает решений самостоятельно — только анализирует и рекомендует",
    "Система не ставит медицинских диагнозов и не назначает медикаменты",
    "Система не даёт инвестиционных советов и не проводит транзакции",
    "Семейные блоки в расписании неприкосновенны",
    "При обнаружении высокого риска в переговорах (флаг high_risk) — сигнал передаётся в Decision Support",
]
for rule in security_rules:
    doc.add_paragraph(rule, style='List Bullet')

# =====================================================================
# SECTION 9 — GLOSSARY
# =====================================================================
add_heading("9. Глоссарий", 1, (88, 56, 186))

glossary = [
    ("LOS", "Life Operating System — персональная мультиагентная ИИ-система"),
    ("CoS", "Chief of Staff — Аня, единственный человеческий интерфейс системы"),
    ("Master Orchestrator", "Центральный агент-координатор, управляющий всеми остальными агентами"),
    ("Readiness Score", "Показатель Oura Ring (0–100): готовность к нагрузке на основе биометрии"),
    ("HRV", "Heart Rate Variability — вариабельность сердечного ритма; ключевой маркер восстановления"),
    ("PLAUDE", "Приложение для записи и транскрипции переговоров"),
    ("CoinKeeper", "Приложение для трекинга личных финансов"),
    ("Mem0", "Open-source фреймворк памяти для ИИ-агентов (Apache 2.0)"),
    ("ReAct", "Reasoning + Acting — паттерн агента: думает → вызывает инструмент → анализирует результат"),
    ("pgvector", "Расширение PostgreSQL для хранения и поиска векторных эмбеддингов"),
    ("high_risk", "Флаг Communication Intelligence Agent при обнаружении высокого риска в переговорах"),
    ("deadline_conflict", "Флаг Decision Support при конфликте финансового дедлайна и других приоритетов"),
]
table10 = doc.add_table(rows=1, cols=2)
table10.style = 'Table Grid'
hdr10 = table10.rows[0].cells
hdr10[0].text = "Термин"
hdr10[1].text = "Определение"
for cell in hdr10:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
for row in glossary:
    add_table_row(table10, row)

# =====================================================================
# FOOTER NOTE
# =====================================================================
doc.add_page_break()
doc.add_paragraph()
footer_note = doc.add_paragraph()
footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_note.add_run("— Конфиденциальный документ. Только для внутреннего использования. —")
run.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.size = Pt(10)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run(f"LOS v1.0  |  {datetime.date.today().strftime('%d.%m.%Y')}")
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.size = Pt(10)

# Save
doc.save('/home/runner/workspace/LOS_Full_Context_Document.docx')
print("Document saved successfully!")
